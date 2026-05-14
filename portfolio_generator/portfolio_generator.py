from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.request
import zipfile
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_FILES = ("index.html", "styles.css", "script.js")
DEFAULT_TARGET_ROLE = "传热 / 热管理 / CFD / 化工工程师"


@dataclass
class ParsedResume:
    text: str
    name: str
    email: str
    phone: str
    city: str
    salary: str
    education: List[Dict[str, str]]
    skills: List[str]
    projects: List[Dict[str, object]]
    certificates: List[str]
    highlights: List[str]


@dataclass
class GenerationResult:
    output_dir: Path
    preview_url: str
    pages_url: Optional[str] = None
    repo_url: Optional[str] = None
    warnings: Optional[List[str]] = None


def generate_site(
    resume_path: Path,
    github_username: str,
    repo_name: str,
    output_dir: Path,
    target_role: str = DEFAULT_TARGET_ROLE,
    city: str = "",
    salary: str = "",
    github_profile: str = "",
    project_doc_path: Optional[Path] = None,
) -> GenerationResult:
    resume_path = Path(resume_path)
    output_dir = Path(output_dir)
    warnings: List[str] = []

    if not resume_path.exists():
        raise FileNotFoundError(f"简历文件不存在：{resume_path}")

    resume_text = extract_pdf_text(resume_path)
    if not resume_text.strip():
        warnings.append("没有从 PDF 中提取到可用文本，网站会使用文件名和表单信息生成基础版本。")

    parsed = parse_resume(resume_text, resume_path.stem)
    if city:
        parsed.city = city
    if salary:
        parsed.salary = salary

    project_doc_text = ""
    if project_doc_path and Path(project_doc_path).exists():
        project_doc_text = extract_docx_text(Path(project_doc_path))
        if not project_doc_text.strip():
            warnings.append("项目文档未提取到文本，问答知识库会仅基于简历生成。")

    data = build_portfolio_data(
        parsed=parsed,
        github_username=github_username.strip(),
        github_profile=github_profile.strip(),
        repo_name=repo_name.strip(),
        target_role=target_role.strip() or DEFAULT_TARGET_ROLE,
        project_doc_text=project_doc_text,
    )

    write_site(output_dir, resume_path, data)
    return GenerationResult(
        output_dir=output_dir,
        preview_url=(output_dir / "index.html").as_uri(),
        warnings=warnings,
    )


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return ""

    chunks: List[str] = []
    try:
        reader = PdfReader(str(path))
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
    except Exception:
        return ""
    return "\n".join(chunks)


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        chunks: List[str] = []
        for paragraph in doc.paragraphs:
            text = normalize_space(paragraph.text)
            if text:
                chunks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [normalize_space(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    chunks.append(row_text)
        return "\n".join(chunks)
    except Exception:
        pass

    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        xml = re.sub(r"<[^>]+>", "", xml)
        return normalize_space(xml.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&"))
    except Exception:
        return ""


def parse_resume(text: str, fallback_name: str) -> ParsedResume:
    text = normalize_resume_text(text)
    lines = [normalize_space(line) for line in text.splitlines()]
    lines = [line for line in lines if line]

    email = first_match(text, r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
    phone = first_match(text, r"(?<!\d)1[3-9]\d{9}(?!\d)")
    name = guess_name(lines, fallback_name)
    city = guess_field(text, "期望城市") or ""
    salary = guess_field(text, "期望薪资") or ""
    education = extract_education(lines)
    skills = extract_skills(text)
    projects = extract_projects(text)
    certificates = extract_certificates(text)
    highlights = extract_highlights(text)

    return ParsedResume(
        text=text,
        name=name,
        email=email,
        phone=phone,
        city=city,
        salary=salary,
        education=education,
        skills=skills,
        projects=projects,
        certificates=certificates,
        highlights=highlights,
    )


def build_portfolio_data(
    parsed: ParsedResume,
    github_username: str,
    github_profile: str,
    repo_name: str,
    target_role: str,
    project_doc_text: str,
) -> Dict[str, object]:
    github_url = github_profile or (f"https://github.com/{github_username}" if github_username else "")
    education_headline = best_education_headline(parsed.education)
    tool_text = ", ".join(parsed.skills[:6]) if parsed.skills else "Python、CFD、ANSYS、Fluent、工程制图"
    project_titles = [str(project.get("title", "")) for project in parsed.projects if project.get("title")]
    core_direction = guess_core_direction(parsed.text, project_titles)

    hero_stats = build_hero_stats(parsed, core_direction)
    facts = [
        {"label": "目标岗位", "value": target_role},
        {"label": "期望城市", "value": parsed.city or "待补充"},
        {"label": "期望薪资", "value": parsed.salary or "待补充"},
        {"label": "核心方向", "value": core_direction},
        {"label": "工程工具", "value": tool_text},
        {"label": "证书", "value": "、".join(parsed.certificates[:3]) if parsed.certificates else "待补充"},
    ]

    return {
        "name": parsed.name,
        "status": f"{target_role} 求职中",
        "headline": education_headline or "传热与工程仿真方向求职者",
        "summary": build_summary(parsed, core_direction),
        "about": build_about(parsed, core_direction),
        "researchIntro": build_research_intro(parsed, core_direction),
        "contactIntro": f"欢迎联系我交流{target_role}相关岗位机会。",
        "resumeNote": "PDF 简历已放入网站资源目录，可直接下载查看完整经历与联系方式。",
        "heroStats": hero_stats,
        "facts": facts,
        "links": {
            "github": github_url,
            "resume": "./assets/resume.pdf",
            "demo": "#demo",
        },
        "contacts": [
            {"label": "邮箱", "value": parsed.email or "待补充邮箱", "href": f"mailto:{parsed.email}" if parsed.email else ""},
            {"label": "电话", "value": parsed.phone or "待补充电话", "href": f"tel:{parsed.phone}" if parsed.phone else ""},
            {"label": "GitHub", "value": github_url.replace("https://", "") if github_url else "待补充 GitHub 链接", "href": github_url},
            {"label": "期望城市", "value": parsed.city or "待补充", "href": ""},
        ],
        "timeline": build_timeline(parsed),
        "researchCards": build_research_cards(parsed, core_direction),
        "publications": build_publications(parsed.text),
        "skills": build_skill_groups(parsed.skills),
        "projects": build_project_cards(parsed, project_doc_text),
        "demos": [
            {"label": "热阻估算器", "href": "#demo"},
            {"label": "项目问答", "href": "#space-qa-title"},
        ],
        "suggestedQuestions": [
            "这个项目主要做什么？",
            "实验模块由哪些部分组成？",
            "它验证了哪些能力？",
            "地面测试有什么结果？",
            "软件系统负责什么？",
            "个人主要参与了什么？",
        ],
        "spaceKnowledge": build_space_knowledge(parsed, project_doc_text),
    }


def write_site(output_dir: Path, resume_path: Path, data: Dict[str, object]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for filename in TEMPLATE_FILES:
        source = ROOT / filename
        if not source.exists():
            raise FileNotFoundError(f"缺少模板文件：{source}")
        shutil.copy2(source, output_dir / filename)

    assets = output_dir / "assets"
    assets.mkdir(exist_ok=True)
    shutil.copy2(resume_path, assets / "resume.pdf")
    (assets / ".gitkeep").write_text("", encoding="utf-8")
    (output_dir / ".nojekyll").write_text("", encoding="utf-8")
    (output_dir / "data.js").write_text(
        "window.portfolioData = " + json.dumps(data, ensure_ascii=False, indent=2) + ";\n",
        encoding="utf-8",
    )
    (output_dir / "README.md").write_text(build_generated_readme(data), encoding="utf-8")


def deploy_to_github_pages(
    site_dir: Path,
    github_username: str,
    repo_name: str,
    user_name: str,
    user_email: str,
    overwrite_remote: bool = False,
) -> Tuple[str, str]:
    token = get_github_token()
    if not token:
        run(["git", "credential-manager", "github", "login"], cwd=site_dir, check=False)
        token = get_github_token()
    if not token:
        raise RuntimeError("没有拿到 GitHub 登录凭据，请先在电脑上登录 GitHub。")

    repo_url, clone_url = create_github_repo(token, repo_name, github_username)
    prepare_git_repo(site_dir, clone_url, user_name, user_email, overwrite_remote)
    pages_url = enable_github_pages(token, github_username, repo_name)
    return repo_url, pages_url


def get_github_token() -> str:
    input_text = "protocol=https\nhost=github.com\n\n"
    commands: List[List[str]] = []
    gcm_path = Path(r"C:\Program Files\Git\mingw64\bin\git-credential-manager.exe")
    if gcm_path.exists():
        commands.append([str(gcm_path), "get"])
    commands.append(["git", "credential-manager", "get"])

    for command in commands:
        try:
            proc = subprocess.run(
                command,
                input=input_text,
                text=True,
                capture_output=True,
                timeout=30,
            )
        except Exception:
            continue
        if proc.returncode != 0:
            continue
        for line in proc.stdout.splitlines():
            if line.startswith("password="):
                return line.split("=", 1)[1].strip()
    return ""


def create_github_repo(token: str, repo_name: str, owner: str) -> Tuple[str, str]:
    headers = github_headers(token)
    existing = request_json(
        f"https://api.github.com/repos/{owner}/{repo_name}",
        headers=headers,
        method="GET",
        allow_404=True,
    )
    if existing:
        return existing["html_url"], existing["clone_url"]

    body = {
        "name": repo_name,
        "description": "Personal portfolio generated from resume",
        "private": False,
        "auto_init": False,
        "has_issues": False,
        "has_projects": False,
        "has_wiki": False,
    }
    created = request_json(
        "https://api.github.com/user/repos",
        headers=headers,
        method="POST",
        data=body,
    )
    return created["html_url"], created["clone_url"]


def enable_github_pages(token: str, owner: str, repo_name: str) -> str:
    headers = github_headers(token)
    body = {"source": {"branch": "main", "path": "/"}}
    try:
        request_json(
            f"https://api.github.com/repos/{owner}/{repo_name}/pages",
            headers=headers,
            method="POST",
            data=body,
        )
    except urllib.error.HTTPError as error:
        if error.code not in (409, 422):
            raise
        request_json(
            f"https://api.github.com/repos/{owner}/{repo_name}/pages",
            headers=headers,
            method="PUT",
            data=body,
        )
    return f"https://{owner}.github.io/{repo_name}/"


def prepare_git_repo(site_dir: Path, clone_url: str, user_name: str, user_email: str, overwrite_remote: bool) -> None:
    if not (site_dir / ".git").exists():
        run(["git", "init"], cwd=site_dir)
    run(["git", "branch", "-M", "main"], cwd=site_dir)
    if user_name:
        run(["git", "config", "user.name", user_name], cwd=site_dir)
    if user_email:
        run(["git", "config", "user.email", user_email], cwd=site_dir)
    if run(["git", "remote", "get-url", "origin"], cwd=site_dir, check=False).returncode == 0:
        run(["git", "remote", "set-url", "origin", clone_url], cwd=site_dir)
    else:
        run(["git", "remote", "add", "origin", clone_url], cwd=site_dir)
    run(["git", "add", "."], cwd=site_dir)
    if run(["git", "diff", "--cached", "--quiet"], cwd=site_dir, check=False).returncode != 0:
        run(["git", "commit", "-m", "Generate portfolio site"], cwd=site_dir)
    push = ["git", "push", "-u", "origin", "main"]
    if overwrite_remote:
        push.insert(2, "--force")
    run(push, cwd=site_dir)


def request_json(
    url: str,
    headers: Dict[str, str],
    method: str = "GET",
    data: Optional[Dict[str, object]] = None,
    allow_404: bool = False,
) -> Dict[str, object]:
    payload = None
    if data is not None:
        payload = json.dumps(data).encode("utf-8")
        headers = {**headers, "Content-Type": "application/json"}
    request = urllib.request.Request(url, data=payload, headers=headers, method=method)
    try:
        with urllib.request.urlopen(request, timeout=60) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except urllib.error.HTTPError as error:
        if allow_404 and error.code == 404:
            return {}
        raise


def github_headers(token: str) -> Dict[str, str]:
    return {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": "portfolio-generator",
    }


def run(command: List[str], cwd: Path, check: bool = True) -> subprocess.CompletedProcess:
    proc = subprocess.run(command, cwd=str(cwd), text=True, capture_output=True)
    if check and proc.returncode != 0:
        raise RuntimeError(
            f"命令执行失败：{' '.join(command)}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def normalize_resume_text(text: str) -> str:
    return text.replace("\ufb02", "fl").replace("\r\n", "\n").replace("\r", "\n")


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def first_match(text: str, pattern: str) -> str:
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def guess_field(text: str, label: str) -> str:
    pattern = rf"{label}[：:]\s*([^\n|]+)"
    match = re.search(pattern, text)
    if match:
        return normalize_space(match.group(1))
    return ""


def guess_name(lines: List[str], fallback_name: str) -> str:
    for line in lines[:10]:
        if re.search(r"@|电话|手机|男|女|\d", line):
            continue
        if 2 <= len(line) <= 12:
            return line
    return fallback_name.replace("简历", "").strip() or "你的姓名"


def extract_education(lines: List[str]) -> List[Dict[str, str]]:
    education: List[Dict[str, str]] = []
    for line in lines:
        if len(line) > 80 or re.match(r"^\d+[.、]", line):
            continue
        if "大学" not in line and "学院" not in line:
            continue
        if not re.search(r"本科|硕士|博士|专科|学士|研究生", line):
            continue
        if not re.search(r"20\d{2}", line) and not re.search(r"\s(本科|硕士|博士|专科|学士|研究生)\s", line):
            continue
        period = first_match(line, r"(20\d{2})\s*[-~—至]\s*(20\d{2}|至今)")
        school = line
        degree = ""
        major = ""
        degree_match = re.search(r"(博士|硕士|本科|专科|学士|研究生)", line)
        if degree_match:
            degree = degree_match.group(1)
            school = line[: degree_match.start()].strip(" |")
        period_match = re.search(r"20\d{2}", line)
        major_area = line[degree_match.end() : period_match.start() if period_match else len(line)].strip(" |") if degree_match else line
        major_match = re.search(r"(化学工程与技术|机械设计/制造|能源与动力工程|工程热物理|热能工程|材料科学与工程|[^\s|]{2,12}工程[^\s|]{0,8})", major_area)
        if major_match:
            major = major_match.group(1)
        education.append({"title": school, "subtitle": f"{degree} | {major}".strip(" |"), "period": period or "时间待补充"})
    return unique_dicts(education, "title")[:4]


def extract_skills(text: str) -> List[str]:
    candidates = [
        "Python",
        "C语言",
        "CFD",
        "ANSYS",
        "Fluent",
        "SolidWorks",
        "CAD",
        "MATLAB",
        "COMSOL",
        "Origin",
        "Excel",
        "热仿真",
        "力学仿真",
        "实验数据分析",
        "技术写作",
    ]
    skills = [item for item in candidates if re.search(re.escape(item), text, flags=re.IGNORECASE)]
    section = section_text(text, "专业技能", ["项目经历", "教育经历", "实习经历", "志愿者经历", "资格证书"])
    for line in section.splitlines():
        line = re.sub(r"^\d+[.、]\s*", "", normalize_space(line))
        if line and len(line) <= 30:
            skills.append(line)
    return unique_strings(skills)[:12]


def extract_projects(text: str) -> List[Dict[str, object]]:
    section = section_text(text, "项目经历", ["资格证书", "专业技能", "志愿者经历", "实习经历", "教育经历"])
    lines = [normalize_space(line) for line in section.splitlines()]
    lines = [line for line in lines if line]
    projects: List[Dict[str, object]] = []
    current: Optional[Dict[str, object]] = None
    pending_period = ""

    for line in lines:
        if re.fullmatch(r"20\d{2}[^一-龥A-Za-z]{0,8}(至今|20\d{2})?", line):
            pending_period = line
            continue
        is_bullet = bool(re.match(r"^\d+[.、]\s*", line))
        looks_title = (
            not is_bullet
            and len(line) >= 8
            and bool(re.search(r"研究|项目|实验|设计|分析|系统|优化|平台|模块", line))
        )
        if looks_title:
            current = {
                "title": re.sub(r"\s+第二负责人.*$", "", line),
                "period": pending_period or "项目经历",
                "tags": infer_project_tags(line),
                "bullets": [],
            }
            if "第二负责人" in line:
                current["tags"].append("第二负责人")
            projects.append(current)
            pending_period = ""
            continue
        if current and is_bullet:
            bullet = re.sub(r"^\d+[.、]\s*", "", line)
            current.setdefault("bullets", []).append(bullet)
    return projects[:5]


def extract_certificates(text: str) -> List[str]:
    known = ["大学英语六级", "普通话二级甲等", "大学英语四级"]
    return [item for item in known if item in text]


def extract_highlights(text: str) -> List[str]:
    section = section_text(text, "个人优势", ["教育经历", "实习经历", "项目经历", "专业技能"])
    highlights: List[str] = []
    for line in section.splitlines():
        line = re.sub(r"^\d+[.、]\s*", "", normalize_space(line))
        if line and 8 <= len(line) <= 80:
            highlights.append(line)
    return highlights[:6]


def section_text(text: str, start: str, end_labels: Iterable[str]) -> str:
    start_index = text.find(start)
    if start_index < 0:
        return ""
    end_index = len(text)
    for label in end_labels:
        index = text.find(label, start_index + len(start))
        if index > start_index:
            end_index = min(end_index, index)
    return text[start_index + len(start) : end_index]


def best_education_headline(education: List[Dict[str, str]]) -> str:
    if not education:
        return ""
    first = education[0]
    subtitle = str(first.get("subtitle", ""))
    return f"{first.get('title', '')}{subtitle.replace(' | ', '')}".strip()


def guess_core_direction(text: str, project_titles: List[str]) -> str:
    if "微重力" in text or any("微重力" in title for title in project_titles):
        return "微重力沸腾传热、强化换热与实验验证"
    if "传热" in text or "CFD" in text.upper():
        return "传热、热管理、CFD仿真与实验验证"
    return "工程仿真、实验分析与项目推进"


def build_hero_stats(parsed: ParsedResume, core_direction: str) -> List[Dict[str, str]]:
    stats: List[Dict[str, str]] = []
    score = first_match(parsed.text, r"考研成绩\s*\d+|\b[34]\d{2}\b")
    if score:
        stats.append({"value": re.sub(r"\D", "", score), "label": "考研成绩"})
    if "前5%" in parsed.text or "前 5%" in parsed.text:
        stats.append({"value": "前 5%", "label": "专业排名"})
    if "International Journal of Heat and Mass Transfer" in parsed.text or "国际传热传质" in parsed.text:
        stats.append({"value": "IJHMT", "label": "国际期刊经历"})
    if "第二负责人" in parsed.text:
        stats.append({"value": "第二负责人", "label": "核心科研项目"})
    defaults = [
        {"value": "PDF", "label": "简历自动解析"},
        {"value": "GitHub", "label": "在线作品集"},
        {"value": "Demo", "label": "内置交互演示"},
        {"value": "Q&A", "label": "项目问答知识库"},
    ]
    for item in defaults:
        if len(stats) >= 4:
            break
        stats.append(item)
    return stats[:4]


def build_summary(parsed: ParsedResume, core_direction: str) -> str:
    return f"聚焦{core_direction}，具备实验平台搭建、数据分析、工程仿真和技术表达能力。"


def build_about(parsed: ParsedResume, core_direction: str) -> str:
    education = best_education_headline(parsed.education)
    base = f"我具备{core_direction}相关背景"
    if education:
        base = f"我目前/曾就读于{education}，具备{core_direction}相关背景"
    if parsed.highlights:
        return base + "。" + parsed.highlights[0]
    return base + "，能够将理论分析、仿真建模、实验验证和结果表达结合起来推进工程问题。"


def build_research_intro(parsed: ParsedResume, core_direction: str) -> str:
    if "未充分发展" in parsed.text:
        return "当前重点研究未充分发展流动中流速对传热性能的影响，已形成初步结果并进入论文撰写阶段。"
    return f"当前关注{core_direction}相关问题，强调实验数据、仿真分析和工程结论之间的相互验证。"


def build_timeline(parsed: ParsedResume) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    for edu in parsed.education:
        items.append(
            {
                "title": edu.get("title", "教育经历"),
                "subtitle": edu.get("subtitle", "专业经历"),
                "period": edu.get("period", "时间待补充"),
                "detail": "系统学习专业课程并参与科研/工程训练，形成理论分析与项目实践基础。",
            }
        )
    if "实习经历" in parsed.text:
        items.append(
            {
                "title": "工程实习经历",
                "subtitle": "现场认知 | 工程实践",
                "period": "实习阶段",
                "detail": "参与企业或工程现场实践，理解安全规范、工艺流程和工程问题的实际约束。",
            }
        )
    if "公众号" in parsed.text:
        items.append(
            {
                "title": "专业内容编辑",
                "subtitle": "行业动态 | 技术表达",
                "period": "学习阶段",
                "detail": "参与专业公众号编辑，持续跟踪行业前沿，提升信息筛选和技术表达能力。",
            }
        )
    if "任教" in parsed.text or "教育" in parsed.text:
        items.append(
            {
                "title": "教学与沟通经历",
                "subtitle": "知识讲授 | 结构化表达",
                "period": "兼职/实践",
                "detail": "通过教学或讲授经历训练复杂知识拆解、沟通反馈和快速学习能力。",
            }
        )
    if not items:
        items.append(
            {
                "title": "教育与项目经历",
                "subtitle": "待补充",
                "period": "待补充",
                "detail": "可在生成后继续补充更详细的学校、项目、实习和成果信息。",
            }
        )
    return items[:6]


def build_research_cards(parsed: ParsedResume, core_direction: str) -> List[Dict[str, str]]:
    return [
        {"title": "当前方向", "body": build_research_intro(parsed, core_direction)},
        {"title": "方法优势", "body": "将实验测量、仿真建模、数据分析和机理解释结合起来，关注结论的可验证性。"},
        {"title": "岗位匹配", "body": "适合需要传热基础、仿真能力、实验分析和跨团队沟通的工程研发岗位。"},
    ]


def build_publications(text: str) -> List[Dict[str, str]]:
    publications: List[Dict[str, str]] = []
    if "Flow-jet composite boiling" in text:
        publications.append(
            {
                "title": "Flow-jet composite boiling on microstructured surfaces in microgravity: an experimental study",
                "venue": "International Journal of Heat and Mass Transfer",
                "role": "第三作者",
                "status": "已发表",
            }
        )
    if "论文" in text:
        publications.append(
            {
                "title": "相关研究论文",
                "venue": "围绕实验、仿真与机理分析展开",
                "role": "参与撰写",
                "status": "进行中",
            }
        )
    return publications or [
        {
            "title": "科研与项目成果",
            "venue": "可在生成后继续补充论文、专利、竞赛或项目链接",
            "role": "待补充",
            "status": "待完善",
        }
    ]


def build_skill_groups(skills: List[str]) -> List[Dict[str, object]]:
    skill_text = " ".join(skills)
    return [
        {
            "title": "传热与流动",
            "items": ["传热学与流体力学基础", "换热过程分析", "热阻网络与温升估算", "实验现象机理解释"],
        },
        {
            "title": "仿真建模",
            "items": [item for item in ["CFD 建模与后处理", "ANSYS / Fluent 模拟", "热仿真与力学仿真", "边界条件与模型搭建"] if item.split()[0] in skill_text or True],
        },
        {
            "title": "实验与测试",
            "items": ["实验方案设计", "数据采集与整理", "设备维护与问题排查", "实验报告与结果解读"],
        },
        {
            "title": "编程与表达",
            "items": ["Python / C 语言", "数据处理与可视化", "技术写作与论文表达", "沟通协作与快速学习"],
        },
    ]


def build_project_cards(parsed: ParsedResume, project_doc_text: str) -> List[Dict[str, object]]:
    cards: List[Dict[str, object]] = []
    for index, project in enumerate(parsed.projects[:4]):
        title = str(project.get("title", "项目经历"))
        bullets = [str(item) for item in project.get("bullets", [])][:4]
        if not bullets:
            bullets = ["梳理项目目标与技术路径", "参与数据分析、实验验证或仿真建模", "沉淀项目报告和结果表达"]
        cards.append(
            {
                "title": title,
                "period": str(project.get("period", "项目经历")),
                "tags": project.get("tags", ["项目经历"]),
                "featured": index == 0,
                "summary": f"围绕{title}开展实验、仿真、数据分析或工程验证工作。",
                "bullets": bullets,
                "links": [{"label": "项目问答", "href": "#space-qa-title"}] if index == 0 else [{"label": "热阻估算 Demo", "href": "#demo"}],
            }
        )
    if project_doc_text and not cards:
        cards.append(
            {
                "title": "项目文档提炼经历",
                "period": "项目经历",
                "tags": ["文档提炼", "项目问答"],
                "featured": True,
                "summary": "基于上传的项目文档生成公开展示版项目介绍和问答知识库。",
                "bullets": extract_snippets(project_doc_text, ["概述", "功能", "测试", "结论"], 4),
                "links": [{"label": "项目问答", "href": "#space-qa-title"}],
            }
        )
    if not cards:
        cards.append(
            {
                "title": "核心项目经历",
                "period": "待补充",
                "tags": ["项目", "能力展示"],
                "featured": True,
                "summary": "可在生成后继续补充项目背景、个人职责、技术方法和结果产出。",
                "bullets": ["说明项目目标和业务/科研背景", "写清个人负责内容", "补充可量化结果和工具方法"],
                "links": [{"label": "项目问答", "href": "#space-qa-title"}],
            }
        )
    return cards[:4]


def build_space_knowledge(parsed: ParsedResume, project_doc_text: str) -> List[Dict[str, object]]:
    entries: List[Dict[str, object]] = []
    source = project_doc_text or parsed.text
    topics = [
        ("项目主要做什么", ["项目", "概述", "目标", "背景", "研究"], ["项目", "做什么", "目标", "背景", "研究"]),
        ("实验模块组成", ["组成", "模块", "单元", "系统"], ["组成", "模块", "单元", "系统"]),
        ("模块功能能力", ["功能", "测量", "控制", "参数", "数据"], ["功能", "能力", "测量", "控制", "数据"]),
        ("研制与验证过程", ["测试", "试验", "验证", "匹配", "验收"], ["测试", "试验", "验证", "匹配", "验收"]),
        ("结果与结论", ["结果", "结论", "满足", "提高", "通过"], ["结果", "结论", "满足", "提高", "通过"]),
    ]
    for title, search_terms, keywords in topics:
        points = extract_snippets(source, search_terms, 3)
        if points:
            entries.append({"title": title, "keywords": keywords, "points": points})
    entries.append(
        {
            "title": "个人参与价值",
            "keywords": ["你做", "个人", "职责", "负责", "参与", "角色"],
            "points": build_personal_points(parsed),
        }
    )
    return entries


def extract_snippets(text: str, terms: List[str], limit: int) -> List[str]:
    sentences = split_sentences(text)
    results: List[str] = []
    for sentence in sentences:
        if any(term in sentence for term in terms) and 16 <= len(sentence) <= 130:
            results.append(sentence)
        if len(results) >= limit:
            break
    return results


def split_sentences(text: str) -> List[str]:
    text = normalize_space(text)
    parts = re.split(r"[。；;!?！？]\s*", text)
    return [part.strip() for part in parts if part.strip()]


def build_personal_points(parsed: ParsedResume) -> List[str]:
    points = []
    if parsed.projects:
        points.append("参与项目推进，覆盖实验设计、数据分析、仿真建模或报告撰写等环节。")
    if parsed.skills:
        points.append("工具与方法基础包括：" + "、".join(parsed.skills[:6]) + "。")
    if parsed.highlights:
        points.append(parsed.highlights[0])
    return points or ["此部分可在生成后继续补充个人职责、技术方法和项目结果。"]


def infer_project_tags(title: str) -> List[str]:
    tags: List[str] = []
    for keyword in ["微重力", "沸腾", "传热", "CFD", "仿真", "实验", "热管理", "化工"]:
        if keyword in title.upper() or keyword in title:
            tags.append(keyword)
    return tags or ["项目经历"]


def unique_strings(items: Iterable[str]) -> List[str]:
    seen = set()
    result = []
    for item in items:
        item = normalize_space(item)
        if not item or item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def unique_dicts(items: Iterable[Dict[str, str]], key: str) -> List[Dict[str, str]]:
    seen = set()
    result = []
    for item in items:
        marker = item.get(key, "")
        if not marker or marker in seen:
            continue
        seen.add(marker)
        result.append(item)
    return result


def build_generated_readme(data: Dict[str, object]) -> str:
    name = data.get("name", "个人")
    return f"""# {name}的求职作品集网站

本网站由 `portfolio_generator` 根据简历自动生成。

## 使用方式

- 直接打开 `index.html` 预览。
- 上传到 GitHub Pages 后，可通过公开网址访问。
- 简历 PDF 位于 `assets/resume.pdf`。

## 修改内容

主要个人信息、项目、技能和问答知识库位于 `data.js`。
"""


def main(argv: Optional[List[str]] = None) -> int:
    import argparse

    parser = argparse.ArgumentParser(description="根据简历生成求职作品集网站")
    parser.add_argument("--resume", required=True, help="简历 PDF 路径")
    parser.add_argument("--github", required=True, help="GitHub 用户名")
    parser.add_argument("--repo", default="portfolio-site", help="GitHub 仓库名 / 输出目录名")
    parser.add_argument("--output", default="", help="输出目录，默认 generated_sites/<repo>")
    parser.add_argument("--role", default=DEFAULT_TARGET_ROLE, help="目标岗位")
    parser.add_argument("--city", default="", help="期望城市")
    parser.add_argument("--salary", default="", help="期望薪资")
    parser.add_argument("--project-doc", default="", help="可选项目 DOCX 文档")
    parser.add_argument("--deploy", action="store_true", help="生成后部署到 GitHub Pages")
    parser.add_argument("--force", action="store_true", help="部署时覆盖远程 main 分支")
    args = parser.parse_args(argv)

    output = Path(args.output) if args.output else ROOT / "generated_sites" / args.repo
    result = generate_site(
        resume_path=Path(args.resume),
        github_username=args.github,
        repo_name=args.repo,
        output_dir=output,
        target_role=args.role,
        city=args.city,
        salary=args.salary,
        project_doc_path=Path(args.project_doc) if args.project_doc else None,
    )
    print(f"已生成：{result.output_dir}")
    print(f"本地预览：{result.preview_url}")
    if result.warnings:
        for warning in result.warnings:
            print(f"提醒：{warning}")
    if args.deploy:
        repo_url, pages_url = deploy_to_github_pages(
            site_dir=result.output_dir,
            github_username=args.github,
            repo_name=args.repo,
            user_name=parse_resume(extract_pdf_text(Path(args.resume)), Path(args.resume).stem).name,
            user_email=parse_resume(extract_pdf_text(Path(args.resume)), Path(args.resume).stem).email,
            overwrite_remote=args.force,
        )
        print(f"GitHub 仓库：{repo_url}")
        print(f"网站地址：{pages_url}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
